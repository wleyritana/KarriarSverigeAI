from docx import Document
from io import BytesIO
from datetime import datetime

def build_report(sections):
    doc = Document()
    doc.add_heading("AI Job Hunt v5 Intelligence Report", level=1)
    doc.add_paragraph(datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"))
    for title, content in sections.items():
        doc.add_heading(title, level=2)
        doc.add_paragraph(content)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
